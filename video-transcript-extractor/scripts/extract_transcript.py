#!/usr/bin/env python3
"""
Multi-Platform Video Transcript Extractor
Supports: YouTube, Bilibili, Douyin, Toutiao, Ixigua, XiaoHongShu, Zhihu,
          WeChat Channels (微信视频号), and generic video pages.

Three-tier transcript extraction:
  1. Captions/subtitles via yt-dlp (fastest, most accurate)
  2. Whisper speech-to-text via faster-whisper (fallback)
  3. Metadata-only document (last resort)
Author: song2024@gmail.com

"""

import asyncio
import os
import re
import subprocess
import sys
import json
from pathlib import Path
from typing import Optional, List
from urllib.parse import urlparse

# ---------------------------------------------------------------------------
# Auto-install helpers
# ---------------------------------------------------------------------------

def _pip_install(*packages):
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "--break-system-packages", *packages],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )

def _ensure(module_name, pip_name=None):
    try:
        return __import__(module_name)
    except ImportError:
        print(f"Installing {pip_name or module_name}...")
        _pip_install(pip_name or module_name)
        return __import__(module_name)

# Core deps
_ensure("docx", "python-docx")
_ensure("langdetect")

import yt_dlp  # assumed present (CLI tool)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langdetect import detect as detect_lang

# Optional deps — loaded lazily
HAS_TRANSLATOR = False
try:
    from googletrans import Translator as GoogleTranslator
    HAS_TRANSLATOR = True
except ImportError:
    pass

# ---------------------------------------------------------------------------
# Platform detection
# ---------------------------------------------------------------------------

PLATFORM_PATTERNS = [
    (r"youtube\.com|youtu\.be", "youtube"),
    (r"bilibili\.com|b23\.tv", "bilibili"),
    (r"douyin\.com", "douyin"),
    (r"toutiao\.com", "toutiao"),
    (r"ixigua\.com", "ixigua"),
    (r"xiaohongshu\.com|xhslink\.com", "xiaohongshu"),
    (r"zhihu\.com", "zhihu"),
    (r"channels\.weixin\.qq\.com|finder\.video\.qq\.com", "wechat_channels"),
    (r"kuaishou\.com", "kuaishou"),
    (r"weibo\.com", "weibo"),
]

PLAYWRIGHT_ONLY_PLATFORMS = {"wechat_channels"}

PLATFORM_REFERERS = {
    "toutiao": "https://www.toutiao.com/",
    "douyin": "https://www.douyin.com/",
    "bilibili": "https://www.bilibili.com/",
    "ixigua": "https://www.ixigua.com/",
    "wechat_channels": "https://channels.weixin.qq.com/",
    "kuaishou": "https://www.kuaishou.com/",
    "xiaohongshu": "https://www.xiaohongshu.com/",
    "weibo": "https://weibo.com/",
}

LANG_NAMES = {
    "en": "English", "zh": "Chinese", "zh-cn": "Chinese", "zh-tw": "Chinese (Traditional)",
    "ja": "Japanese", "ko": "Korean", "es": "Spanish", "fr": "French", "de": "German",
    "ru": "Russian", "pt": "Portuguese", "ar": "Arabic",
}

ALL_TRANSCRIPT_FORMATS = ["docx", "txt", "pdf", "md"]
ALL_VIDEO_FORMATS = ["mp4", "mp3", "webm", "mkv"]


def detect_platform(url: str) -> str:
    for pattern, name in PLATFORM_PATTERNS:
        if re.search(pattern, url, re.I):
            return name
    return "generic"


def sanitize_filename(filename: str) -> str:
    return re.sub(r'[<>:"/\\|?*\n\r]', '', filename).strip()


def _is_chinese(lang: str) -> bool:
    return lang.lower() in ("zh", "zh-cn", "zh-tw", "zh-hans", "zh-hant")


def _group_sentences(text: str, group_size: int = 5) -> List[str]:
    """Group short Whisper segments into logical paragraphs."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) <= group_size:
        return [" ".join(lines) if not _is_cjk_text(" ".join(lines)) else "".join(lines)]
    paragraphs = []
    for i in range(0, len(lines), group_size):
        chunk = lines[i:i + group_size]
        joined = "".join(chunk) if _is_cjk_text(chunk[0]) else " ".join(chunk)
        paragraphs.append(joined)
    return paragraphs


def _is_cjk_text(text: str) -> bool:
    """Check if text is predominantly CJK characters."""
    if not text:
        return False
    cjk = sum(1 for c in text if '一' <= c <= '鿿' or '　' <= c <= '〿')
    return cjk / max(len(text), 1) > 0.3


# ---------------------------------------------------------------------------
# Video download — Strategy 1: yt-dlp
# ---------------------------------------------------------------------------

def get_video_info_ytdlp(url: str) -> Optional[dict]:
    print(f"📥 Extracting video info from {url}...")
    try:
        opts = {"quiet": True, "no_warnings": True}
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(url, download=False)
        return {
            "title": info.get("title", "Unknown"),
            "description": info.get("description", ""),
            "duration": info.get("duration", 0),
            "channel": info.get("channel") or info.get("uploader", "Unknown"),
            "upload_date": info.get("upload_date", "Unknown"),
        }
    except Exception as e:
        print(f"⚠️ yt-dlp info extraction failed: {e}")
        return None


def download_video_ytdlp(url: str, output_dir: str, platform: str,
                         video_formats: List[str] = None) -> List[str]:
    """Download video in requested formats. Returns list of downloaded file paths."""
    if video_formats is None:
        video_formats = ["mp4"]

    downloaded_files = []

    for fmt in video_formats:
        print(f"📹 Downloading {'audio' if fmt == 'mp3' else 'video'} as {fmt}...")

        if fmt == "mp3":
            opts = {
                "format": "bestaudio/best",
                "outtmpl": os.path.join(output_dir, "%(title).100B[%(id)s].%(ext)s"),
                "quiet": False,
                "no_warnings": True,
                "postprocessors": [{
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "mp3",
                    "preferredquality": "192",
                }],
            }
        else:
            merge_fmt = fmt if fmt in ("mp4", "webm", "mkv") else "mp4"
            opts = {
                "format": "bestvideo+bestaudio/best",
                "outtmpl": os.path.join(output_dir, "%(title).100B[%(id)s].%(ext)s"),
                "quiet": False,
                "no_warnings": True,
                "merge_output_format": merge_fmt,
            }

        if platform == "bilibili":
            opts["referer"] = "https://www.bilibili.com/"

        try:
            with yt_dlp.YoutubeDL(opts) as ydl:
                info = ydl.extract_info(url, download=True)
                video_file = ydl.prepare_filename(info)
                # For mp3, yt-dlp changes extension after postprocessing
                if fmt == "mp3":
                    base = os.path.splitext(video_file)[0]
                    if os.path.exists(base + ".mp3"):
                        video_file = base + ".mp3"
                elif not video_file.endswith(f".{fmt}"):
                    base = os.path.splitext(video_file)[0]
                    if os.path.exists(base + f".{fmt}"):
                        video_file = base + f".{fmt}"
                print(f"✅ Downloaded: {os.path.basename(video_file)}")
                downloaded_files.append(video_file)
        except Exception as e:
            print(f"⚠️ yt-dlp download ({fmt}) failed: {e}")

    return downloaded_files


def extract_captions_ytdlp(url: str) -> Optional[str]:
    print("🔤 Attempting to extract captions via yt-dlp...")
    tmp_prefix = "/tmp/_vte_captions"
    try:
        opts = {
            "skip_download": True,
            "writesubtitles": True,
            "writeautomaticsub": True,
            "subtitleslangs": ["zh-Hans", "zh", "zh-CN", "en", "all"],
            "quiet": True,
            "no_warnings": True,
            "outtmpl": tmp_prefix,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            ydl.extract_info(url, download=True)

        sub_files = sorted(Path("/tmp").glob("_vte_captions*.vtt"))
        if not sub_files:
            sub_files = sorted(Path("/tmp").glob("_vte_captions*.srt"))
        if not sub_files:
            print("⚠️ No caption files found")
            return None

        best = sub_files[0]
        with open(best, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.split("\n")
        text_lines = []
        seen = set()
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("WEBVTT") or stripped.startswith("Kind:") or stripped.startswith("Language:"):
                continue
            if "-->" in stripped:
                continue
            if re.match(r"^\d+$", stripped):
                continue
            cleaned = re.sub(r"<[^>]+>", "", stripped)
            if cleaned and cleaned not in seen:
                seen.add(cleaned)
                text_lines.append(cleaned)

        transcript = " ".join(text_lines)
        if transcript.strip():
            print(f"✅ Extracted {len(transcript)} characters from captions")
            return transcript

    except Exception as e:
        print(f"⚠️ Caption extraction failed: {e}")
    finally:
        for f in Path("/tmp").glob("_vte_captions*"):
            try:
                f.unlink()
            except OSError:
                pass
    return None


# ---------------------------------------------------------------------------
# Video download — Strategy 2: Playwright (for WeChat Channels & fallback)
# ---------------------------------------------------------------------------

async def download_video_playwright(url: str, output_dir: str) -> Optional[str]:
    print("📹 Downloading video via Playwright browser automation...")
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        print("Installing playwright...")
        _pip_install("playwright")
        subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"],
                       capture_output=True)
        from playwright.async_api import async_playwright

    video_url = None
    page_title = "video"

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        )
        try:
            print(f"🌐 Opening page: {url}")
            await page.goto(url, wait_until="domcontentloaded", timeout=30000)
            await asyncio.sleep(5)

            page_title = (await page.title()) or "video"

            video_el = await page.query_selector("video")
            if video_el:
                video_url = await video_el.get_attribute("src")
                if video_url and video_url.startswith("//"):
                    video_url = "https:" + video_url

            if not video_url:
                content = await page.content()
                urls = re.findall(r'https?://[^"\'<>\s]+\.mp4[^"\'<>\s]*', content)
                if urls:
                    video_url = urls[0]

        except Exception as e:
            print(f"⚠️ Playwright page load failed: {e}")
        finally:
            await browser.close()

    if not video_url:
        print("❌ Could not find video URL on page")
        return None

    print(f"✅ Found video URL: {video_url[:80]}...")

    import requests
    platform = detect_platform(url)
    referer = PLATFORM_REFERERS.get(platform, url)
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Referer": referer,
    }

    safe_title = sanitize_filename(page_title)[:80] or "video"
    output_path = os.path.join(output_dir, f"{safe_title}.mp4")

    print("⬇️ Downloading video file...")
    resp = requests.get(video_url, headers=headers, stream=True, timeout=120)
    resp.raise_for_status()

    total = int(resp.headers.get("Content-Length", 0))
    downloaded = 0
    with open(output_path, "wb") as f:
        for chunk in resp.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)
                downloaded += len(chunk)
                if total > 0:
                    pct = downloaded / total * 100
                    print(f"\r   Progress: {pct:.1f}% ({downloaded/1024/1024:.1f} MB)", end="")
    print()

    print(f"✅ Video saved: {os.path.basename(output_path)}")
    return output_path


# ---------------------------------------------------------------------------
# Whisper speech-to-text (Tier 2 fallback)
# ---------------------------------------------------------------------------

def transcribe_with_whisper(video_path: str, language: str = None,
                            model_size: str = "base") -> Optional[str]:
    print("🎤 Transcribing audio with Whisper (this may take a few minutes)...")

    audio_path = video_path + ".wav"
    cmd = [
        "ffmpeg", "-y", "-i", video_path,
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        audio_path,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"⚠️ ffmpeg audio extraction failed: {result.stderr[-200:]}")
        return None

    try:
        try:
            from faster_whisper import WhisperModel
        except ImportError:
            print("Installing faster-whisper...")
            _pip_install("faster-whisper")
            from faster_whisper import WhisperModel

        print(f"🤖 Loading Whisper model ({model_size})...")
        model = WhisperModel(model_size, device="cpu", compute_type="int8")

        kwargs = {"beam_size": 5}
        if language:
            kwargs["language"] = language

        print("📝 Transcribing...")
        segments, info = model.transcribe(audio_path, **kwargs)

        detected = info.language
        print(f"🔤 Detected language: {detected} (probability: {info.language_probability:.2f})")

        texts = [seg.text.strip() for seg in segments]
        transcript = "\n".join(texts)
        print(f"✅ Transcription complete: {len(texts)} segments, {len(transcript)} characters")
        return transcript

    except Exception as e:
        print(f"⚠️ Whisper transcription failed: {e}")
        return None
    finally:
        if os.path.exists(audio_path):
            os.remove(audio_path)


# ---------------------------------------------------------------------------
# Language & translation
# ---------------------------------------------------------------------------

def detect_language(text: str) -> str:
    try:
        return detect_lang(text[:500])
    except Exception:
        return "en"


def translate_to_chinese(text: str) -> Optional[str]:
    if not HAS_TRANSLATOR:
        return None
    print("🌐 Translating to Chinese...")
    try:
        translator = GoogleTranslator()
        chunk_size = 3000
        parts = []
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            result = translator.translate(chunk, src_language="auto", target_language="zh-CN")
            parts.append(result["translatedText"] if isinstance(result, dict) else str(result))
        return " ".join(parts)
    except Exception as e:
        print(f"⚠️ Translation failed: {e}")
        return None


# ---------------------------------------------------------------------------
# Document generation — DOCX (enhanced formatting)
# ---------------------------------------------------------------------------

def create_docx(text: str, title: str, language: str, output_dir: str = ".") -> str:
    safe_title = sanitize_filename(title)[:100]
    lang_code = language.split("-")[0] if "-" in language else language
    filename = os.path.join(output_dir, f"{safe_title}-{lang_code}.docx")

    doc = Document()
    is_cjk = _is_chinese(language) or _is_cjk_text(text[:100])

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei" if is_cjk else "Calibri"
    font.size = Pt(11)
    # Set East Asian font for CJK
    if is_cjk:
        from docx.oxml.ns import qn
        rpr = style.element.get_or_add_rPr()
        ea_font = rpr.makeelement(qn("w:rFonts"), {})
        ea_font.set(qn("w:eastAsia"), "Microsoft YaHei")
        rpr.append(ea_font)

    # -- Title --
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(title)
    run.font.size = Pt(18)

    # -- Metadata line --
    lang_name = LANG_NAMES.get(language.lower(), language.upper())
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    run = meta.add_run(f"Language: {lang_name}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # -- Separator --
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(16)
    run = sep.add_run("— — —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # -- Body paragraphs --
    paragraphs = _group_sentences(text)
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(8)
        if is_cjk:
            p.paragraph_format.first_line_indent = Pt(22)
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei" if is_cjk else "Calibri"

    # -- Footer separator --
    doc.add_paragraph()
    footer_sep = doc.add_paragraph()
    footer_sep.paragraph_format.space_before = Pt(16)
    footer_sep.paragraph_format.space_after = Pt(4)
    run = footer_sep.add_run("─" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # -- Footer info --
    footer = doc.add_paragraph()
    run = footer.add_run("Generated by Video Transcript Extractor")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(filename)
    print(f"✅ Saved {os.path.basename(filename)}")
    return filename


# ---------------------------------------------------------------------------
# Document generation — TXT
# ---------------------------------------------------------------------------

def create_txt(text: str, title: str, language: str, output_dir: str = ".") -> str:
    safe_title = sanitize_filename(title)[:100]
    lang_code = language.split("-")[0] if "-" in language else language
    filename = os.path.join(output_dir, f"{safe_title}-{lang_code}.txt")

    lang_name = LANG_NAMES.get(language.lower(), language.upper())
    paragraphs = _group_sentences(text)

    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write(f"{'=' * min(len(title) * 2, 60)}\n")
        f.write(f"Language: {lang_name}\n\n")
        f.write("---\n\n")
        for para in paragraphs:
            if para.strip():
                f.write(f"{para}\n\n")
        f.write("---\n")
        f.write("Generated by Video Transcript Extractor\n")

    print(f"✅ Saved {os.path.basename(filename)}")
    return filename


# ---------------------------------------------------------------------------
# Document generation — Markdown
# ---------------------------------------------------------------------------

def create_markdown(text: str, title: str, language: str, output_dir: str = ".") -> str:
    safe_title = sanitize_filename(title)[:100]
    lang_code = language.split("-")[0] if "-" in language else language
    filename = os.path.join(output_dir, f"{safe_title}-{lang_code}.md")

    lang_name = LANG_NAMES.get(language.lower(), language.upper())
    paragraphs = _group_sentences(text)

    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write(f"> **Language:** {lang_name}\n\n")
        f.write("---\n\n")
        for para in paragraphs:
            if para.strip():
                f.write(f"{para}\n\n")
        f.write("---\n\n")
        f.write("*Generated by Video Transcript Extractor*\n")

    print(f"✅ Saved {os.path.basename(filename)}")
    return filename


# ---------------------------------------------------------------------------
# Document generation — PDF
# ---------------------------------------------------------------------------

def create_pdf(text: str, title: str, language: str, output_dir: str = ".") -> str:
    try:
        from fpdf import FPDF
    except ImportError:
        print("Installing fpdf2...")
        _pip_install("fpdf2")
        from fpdf import FPDF

    safe_title = sanitize_filename(title)[:100]
    lang_code = language.split("-")[0] if "-" in language else language
    filename = os.path.join(output_dir, f"{safe_title}-{lang_code}.pdf")

    is_cjk = _is_chinese(language) or _is_cjk_text(text[:100])
    lang_name = LANG_NAMES.get(language.lower(), language.upper())
    paragraphs = _group_sentences(text)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Try to add CJK font support
    cjk_font_loaded = False
    if is_cjk:
        cjk_font_paths = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "C:/Windows/Fonts/msyh.ttc",
        ]
        for fpath in cjk_font_paths:
            if os.path.exists(fpath):
                try:
                    pdf.add_font("CJK", "", fpath, uni=True)
                    cjk_font_loaded = True
                    break
                except Exception:
                    continue

    body_font = "CJK" if cjk_font_loaded else "Helvetica"

    # Title
    if cjk_font_loaded:
        pdf.set_font("CJK", size=18)
    else:
        pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Metadata
    if cjk_font_loaded:
        pdf.set_font("CJK", size=9)
    else:
        pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(136, 136, 136)
    pdf.cell(0, 6, f"Language: {lang_name}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Separator
    pdf.set_draw_color(204, 204, 204)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(8)

    # Body
    pdf.set_text_color(33, 33, 33)
    if cjk_font_loaded:
        pdf.set_font("CJK", size=11)
    else:
        pdf.set_font("Helvetica", size=11)

    for para in paragraphs:
        if para.strip():
            pdf.multi_cell(0, 7, f"    {para}" if is_cjk else para)
            pdf.ln(4)

    # Footer
    pdf.ln(8)
    pdf.set_draw_color(204, 204, 204)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)
    if cjk_font_loaded:
        pdf.set_font("CJK", size=8)
    else:
        pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(170, 170, 170)
    pdf.cell(0, 5, "Generated by Video Transcript Extractor", new_x="LMARGIN", new_y="NEXT")

    pdf.output(filename)
    print(f"✅ Saved {os.path.basename(filename)}")
    return filename


# ---------------------------------------------------------------------------
# Metadata fallback document
# ---------------------------------------------------------------------------

def generate_metadata_doc(video_info: dict) -> str:
    desc = video_info.get("description", "")
    if desc:
        desc = re.sub(r"https?://\S+", "", desc).strip()
    return (
        f"Video: {video_info.get('title', 'Untitled')}\n"
        f"Channel: {video_info.get('channel', 'Unknown')}\n"
        f"Duration: {int(video_info.get('duration', 0) // 60)}:"
        f"{int(video_info.get('duration', 0) % 60):02d}\n\n"
        f"Description:\n{desc or 'No description available'}\n\n"
        f"Note: This video has no captions and Whisper transcription was not available. "
        f"Content extracted from metadata only."
    )


# ---------------------------------------------------------------------------
# Transcript output dispatcher
# ---------------------------------------------------------------------------

FORMAT_CREATORS = {
    "docx": create_docx,
    "txt": create_txt,
    "md": create_markdown,
    "pdf": create_pdf,
}


def create_transcript_files(text: str, title: str, language: str,
                            output_dir: str, formats: List[str]) -> List[str]:
    """Generate transcript in all requested formats."""
    files = []
    for fmt in formats:
        creator = FORMAT_CREATORS.get(fmt)
        if creator:
            try:
                f = creator(text, title, language, output_dir)
                files.append(f)
            except Exception as e:
                print(f"⚠️ Failed to create {fmt}: {e}")
    return files


# ---------------------------------------------------------------------------
# Main orchestration
# ---------------------------------------------------------------------------

def extract_video_transcript(url: str, output_dir: Optional[str] = None,
                             whisper_model: str = "base",
                             whisper_language: str = None,
                             transcript_formats: List[str] = None,
                             video_formats: List[str] = None) -> dict:
    if output_dir is None:
        output_dir = os.getcwd()
    os.makedirs(output_dir, exist_ok=True)

    if transcript_formats is None:
        transcript_formats = ["docx"]
    if video_formats is None:
        video_formats = ["mp4"]

    platform = detect_platform(url)
    print(f"\n🎬 Processing: {url}")
    print(f"📌 Platform: {platform}")
    print(f"📄 Transcript formats: {', '.join(transcript_formats)}")
    print(f"🎥 Video formats: {', '.join(video_formats)}\n")

    # --- Get video info ---
    video_info = None
    if platform not in PLAYWRIGHT_ONLY_PLATFORMS:
        video_info = get_video_info_ytdlp(url)
    title = (video_info or {}).get("title", "video")

    # --- Download video ---
    video_files = []
    primary_video = None

    if platform in PLAYWRIGHT_ONLY_PLATFORMS:
        pw_file = asyncio.run(download_video_playwright(url, output_dir))
        if pw_file:
            video_files.append(pw_file)
            primary_video = pw_file
    else:
        video_files = download_video_ytdlp(url, output_dir, platform, video_formats)
        if video_files:
            primary_video = video_files[0]
        else:
            print("⚠️ yt-dlp failed, trying Playwright fallback...")
            pw_file = asyncio.run(download_video_playwright(url, output_dir))
            if pw_file:
                video_files.append(pw_file)
                primary_video = pw_file

    if not primary_video or not os.path.exists(primary_video):
        print("❌ Failed to download video")
        return {"success": False, "error": "Video download failed"}

    if title == "video":
        title = Path(primary_video).stem

    print(f"\n📌 Title: {title}\n")

    # --- Tier 1: Try captions ---
    transcript = None
    transcript_source = None

    if platform not in PLAYWRIGHT_ONLY_PLATFORMS:
        transcript = extract_captions_ytdlp(url)
        if transcript:
            transcript_source = "captions"

    # --- Tier 2: Whisper fallback ---
    if not transcript:
        print("\n⚠️ No captions available. Falling back to Whisper transcription...\n")
        # Use primary video (prefer mp4 for audio extraction)
        whisper_input = primary_video
        for vf in video_files:
            if vf.endswith(".mp4"):
                whisper_input = vf
                break
        transcript = transcribe_with_whisper(
            whisper_input,
            language=whisper_language,
            model_size=whisper_model,
        )
        if transcript:
            transcript_source = "whisper"

    # --- Generate transcript documents ---
    output_files = []

    if transcript:
        language = detect_language(transcript)
        print(f"\n🔤 Detected language: {language}")

        original_files = create_transcript_files(
            transcript, title, language, output_dir, transcript_formats)
        output_files.extend(original_files)

        if language not in ("zh", "zh-cn", "zh-tw"):
            translated = translate_to_chinese(transcript)
            if translated:
                zh_files = create_transcript_files(
                    translated, title, "zh", output_dir, transcript_formats)
                output_files.extend(zh_files)

    else:
        print("\n⚠️ Whisper also unavailable. Creating metadata-only document.\n")
        if not video_info:
            video_info = {"title": title, "description": "", "duration": 0, "channel": "Unknown"}
        summary = generate_metadata_doc(video_info)
        fallback_files = create_transcript_files(
            summary, title, "en", output_dir, transcript_formats)
        output_files.extend(fallback_files)
        transcript_source = "metadata"

    # --- Summary ---
    print(f"\n{'=' * 50}")
    print("✅ Processing complete!")
    print(f"{'=' * 50}")
    for vf in video_files:
        print(f"  🎥 {os.path.basename(vf)}")
    for f in output_files:
        print(f"  📄 {os.path.basename(f)}")
    print(f"  📊 Transcript source: {transcript_source}")
    print(f"{'=' * 50}\n")

    return {
        "success": True,
        "title": title,
        "platform": platform,
        "files": video_files + output_files,
        "video_files": video_files,
        "transcript_files": output_files,
        "transcript_source": transcript_source,
    }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_transcript.py <VIDEO_URL> [OUTPUT_DIR]")
        print("       [--model base] [--lang zh]")
        print("       [--transcript-format docx,txt,pdf,md]")
        print("       [--video-format mp4,mp3,webm,mkv]")
        sys.exit(1)

    video_url = sys.argv[1]
    out_dir = None
    model = "base"
    lang = None
    t_fmts = ["docx"]
    v_fmts = ["mp4"]

    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--model" and i + 1 < len(args):
            model = args[i + 1]
            i += 2
        elif args[i] == "--lang" and i + 1 < len(args):
            lang = args[i + 1]
            i += 2
        elif args[i] == "--transcript-format" and i + 1 < len(args):
            t_fmts = [f.strip() for f in args[i + 1].split(",") if f.strip() in ALL_TRANSCRIPT_FORMATS]
            i += 2
        elif args[i] == "--video-format" and i + 1 < len(args):
            v_fmts = [f.strip() for f in args[i + 1].split(",") if f.strip() in ALL_VIDEO_FORMATS]
            i += 2
        else:
            out_dir = args[i]
            i += 1

    result = extract_video_transcript(
        video_url, out_dir,
        whisper_model=model, whisper_language=lang,
        transcript_formats=t_fmts, video_formats=v_fmts,
    )
    sys.exit(0 if result["success"] else 1)
